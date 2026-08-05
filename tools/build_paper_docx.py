"""Build a clean Word manuscript from the current LaTeX source."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TEX = ROOT / "paper" / "main.tex"
DEFAULT_BBL = ROOT / "paper" / "main.bbl"
DEFAULT_ARCHITECTURE = ROOT / "paper" / "figures" / "fig1_rgd_architecture.png"
DEFAULT_FIGURE = ROOT / "paper" / "figures" / "fig2_main_and_ablation.png"
DEFAULT_OUTPUT = ROOT / "paper" / "RGD_manuscript.docx"


LATEX_REPLACEMENTS = {
    r"\rgd": "RGD",
    r"\pi_f": "pi_f",
    r"\widehat\ell_t": "ell_hat_t",
    r"\ell": "ell",
    r"\gamma": "gamma",
    r"\epsilon": "epsilon",
    r"\kappa": "kappa",
    r"\lambda_L": "lambda_L",
    r"\lambda_A": "lambda_A",
    r"\lambda_H": "lambda_H",
    r"\lambda_G": "lambda_G",
    r"\delta_c": "delta_c",
    r"\ge": ">=",
    r"\le": "<=",
    r"\ne": "!=",
    r"\ldots": "...",
    r"\mathrm": "",
    r"\text": "",
    r"\mathbb": "",
    r"\mathcal": "",
    r"\U": "U",
    r"\A": "A",
}


LABEL_NUMBERS = {
    "eq:release_transition": "(1)",
    "eq:matched_return": "(2)",
    "eq:corrective_set": "(3)",
    "eq:maneuver_breadth": "(4)",
    "eq:headroom_need": "(5)",
    "eq:serial_gate": "(6)",
    "eq:release_contract": "(7)",
    "tab:main_summary": "I",
    "fig:architecture": "1",
    "fig:evidence": "2",
}


def strip_tex(text: str, citation_numbers: dict[str, int] | None = None) -> str:
    text = re.sub(r"(?<!\\)%.*", "", text)
    text = text.replace("~", " ").replace("--", "-")
    citation_numbers = citation_numbers or {}

    def replace_citation(match: re.Match[str]) -> str:
        keys = [key.strip() for key in match.group(1).split(",") if key.strip()]
        numbers = []
        for key in keys:
            if key not in citation_numbers:
                raise ValueError(f"citation key missing from .bbl: {key}")
            numbers.append(f"[{citation_numbers[key]}]")
        return ", ".join(numbers)

    text = re.sub(r"\\cite\{([^}]*)\}", replace_citation, text)
    text = re.sub(
        r"\\(?:ref|eqref)\{([^}]*)\}",
        lambda match: LABEL_NUMBERS.get(match.group(1), f"[{match.group(1)}]"),
        text,
    )
    for source, target in LATEX_REPLACEMENTS.items():
        text = text.replace(source, target)
    text = re.sub(r"\\(?:emph|texttt|textbf|mathrm|operatorname)\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+(?:\[[^]]*\])?", "", text)
    text = text.replace("{", "").replace("}", "")
    text = text.replace(r"\,", " ").replace(r"\;", " ").replace(r"\!", "")
    text = text.replace("&", " ").replace("\\", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def equation_text(lines: list[str], citation_numbers: dict[str, int] | None = None) -> str:
    text = " ".join(lines)
    text = re.sub(r"\\label\{[^}]*\}", "", text)
    text = re.sub(r"\\begin\{[^}]+\}|\\end\{[^}]+\}", "", text)
    return strip_tex(text, citation_numbers)


def parse_bibliography(path: Path) -> tuple[list[str], dict[str, int]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    entries = []
    citation_numbers: dict[str, int] = {}
    pattern = re.compile(r"\\bibitem\{([^}]+)\}(.*?)(?=\\bibitem|\Z)", re.S)
    for index, match in enumerate(pattern.finditer(text), 1):
        citation_numbers[match.group(1)] = index
        block = match.group(2)
        cleaned = strip_tex(block)
        if cleaned:
            entries.append(cleaned)
    return entries, citation_numbers


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    tc_pr.append(shading)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def add_main_table(document: Document) -> None:
    rows = [
        ("RGD", "143", "605.92", "29.01", "0.100"),
        ("Fast-only", "0", "605.86", "2.75", "0.000"),
        ("Always-trigger Slow", "180", "599.98", "35.16", "0.044"),
        ("Random trigger", "133", "601.31", "28.01", "0.074"),
        ("Uncertainty trigger", "3", "605.83", "3.11", "0.034"),
        ("Risk trigger", "98", "605.15", "20.45", "0.189"),
    ]
    headers = ("Method", "Queries", "Distance (m)", "Runtime (ms/f)", "Utility")
    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [2600, 1050, 1850, 1900, 1300]
    for index, (cell, header, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
        cell.text = header
        set_cell_width(cell, width)
        set_cell_shading(cell, "D9EAF2")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(9)
        if index:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_values in rows:
        row = table.add_row()
        for index, (cell, value, width) in enumerate(zip(row.cells, row_values, widths)):
            cell.text = value
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            for run in cell.paragraphs[0].runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)
                if row_values[0] == "RGD":
                    run.bold = True
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.add_run(
        "Table I. Closed-loop endpoint and exposure summary over 30 paired seeds."
    ).italic = True


def add_figure(document: Document, figure: Path, caption_text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(figure), width=Inches(6.45))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(caption_text)
    run.italic = True
    run.font.size = Pt(9)


def build(
    tex_path: Path,
    bbl_path: Path,
    architecture_path: Path,
    figure_path: Path,
    output_path: Path,
) -> None:
    lines = tex_path.read_text(encoding="utf-8").splitlines()
    bibliography_entries, citation_numbers = parse_bibliography(bbl_path)
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05
    for style_name, size in (("Title", 17), ("Heading 1", 12), ("Heading 2", 10.5)):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True

    title_match = re.search(r"\\title\{([^}]*)\}", "\n".join(lines))
    title = strip_tex(title_match.group(1)) if title_match else "RGD Manuscript"
    paragraph = document.add_paragraph(title, style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author = document.add_paragraph("Anonymous Authors")
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER

    in_document = False
    in_abstract = False
    in_keywords = False
    in_equation = False
    in_table = False
    in_figure = False
    figure_index = 0
    equation_lines: list[str] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            text = strip_tex(" ".join(paragraph_lines), citation_numbers)
            if text:
                document.add_paragraph(text)
        paragraph_lines = []

    for raw in lines:
        line = raw.strip()
        if line == r"\begin{document}":
            in_document = True
            continue
        if not in_document or line in {r"\maketitle", r"\end{document}"}:
            continue
        if line == r"\begin{abstract}":
            flush_paragraph()
            document.add_heading("Abstract", level=1)
            in_abstract = True
            continue
        if line == r"\end{abstract}":
            flush_paragraph()
            in_abstract = False
            continue
        if line == r"\begin{IEEEkeywords}":
            flush_paragraph()
            in_keywords = True
            paragraph_lines.append("Index Terms-")
            continue
        if line == r"\end{IEEEkeywords}":
            flush_paragraph()
            in_keywords = False
            continue
        if line.startswith(r"\begin{equation}"):
            flush_paragraph()
            in_equation = True
            equation_lines = []
            continue
        if line.startswith(r"\end{equation}"):
            text = equation_text(equation_lines, citation_numbers)
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(text)
            run.font.name = "Cambria Math"
            run.font.size = Pt(10)
            in_equation = False
            continue
        if in_equation:
            equation_lines.append(line)
            continue
        if line.startswith(r"\begin{table}"):
            flush_paragraph()
            in_table = True
            continue
        if line.startswith(r"\end{table}"):
            add_main_table(document)
            in_table = False
            continue
        if in_table:
            continue
        if line.startswith(r"\begin{figure"):
            flush_paragraph()
            in_figure = True
            continue
        if line.startswith(r"\end{figure"):
            figure_index += 1
            if figure_index == 1:
                add_figure(
                    document,
                    architecture_path,
                    "Figure 1. RGD online contract and offline matched evaluator. "
                    "The offline corrective label is not supplied to the query gate.",
                )
            else:
                add_figure(
                    document,
                    figure_path,
                    "Figure 2. Paired main results and component-selectivity audit. "
                    "Panels (a)-(c) use paired-seed intervals; panel (d) uses "
                    "simulator-seed-cluster intervals.",
                )
            in_figure = False
            continue
        if in_figure:
            continue
        if line.startswith(r"\section*{") or line.startswith(r"\section{"):
            flush_paragraph()
            match = re.search(r"\\section\*?\{([^}]*)\}", line)
            if match:
                document.add_heading(strip_tex(match.group(1)), level=1)
            continue
        if line.startswith(r"\subsection{"):
            flush_paragraph()
            match = re.search(r"\\subsection\{([^}]*)\}", line)
            if match:
                document.add_heading(strip_tex(match.group(1)), level=2)
            continue
        if line.startswith((r"\label", r"\bibliographystyle", r"\bibliography")):
            continue
        if line.startswith(r"\begin{enumerate}"):
            flush_paragraph()
            continue
        if line.startswith(r"\end{enumerate}"):
            flush_paragraph()
            continue
        if line.startswith(r"\item"):
            flush_paragraph()
            text = strip_tex(line[len(r"\item") :], citation_numbers)
            document.add_paragraph(text, style="List Number")
            continue
        if not line:
            flush_paragraph()
            continue
        paragraph_lines.append(line)

    flush_paragraph()
    document.add_heading("References", level=1)
    for index, entry in enumerate(bibliography_entries, 1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.add_run(f"[{index}] ").bold = True
        paragraph.add_run(entry)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Anonymous manuscript for review")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--tex", type=Path, default=DEFAULT_TEX)
    parser.add_argument("--bbl", type=Path, default=DEFAULT_BBL)
    parser.add_argument("--architecture", type=Path, default=DEFAULT_ARCHITECTURE)
    parser.add_argument("--figure", type=Path, default=DEFAULT_FIGURE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build(args.tex, args.bbl, args.architecture, args.figure, args.output)
    print(args.output.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
